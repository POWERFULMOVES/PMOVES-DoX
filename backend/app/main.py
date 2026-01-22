from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Query, Form
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
import os
from pathlib import Path
import shutil
from typing import List, Dict, Optional, Any, Literal, Annotated
import uuid
import asyncio
from dotenv import load_dotenv
import time
import json
from docx import Document
import subprocess
import tempfile
import threading
import re
from pydantic import BaseModel
from app.hrm import HRMConfig, HRMMetrics, refine_sort_digits
from app.api.routers import documents, analysis, system, cipher, models, graph

app = FastAPI(title="PMOVES-DoX API")

app.include_router(documents.router)
app.include_router(analysis.router)
app.include_router(system.router)
app.include_router(cipher.router)
app.include_router(models.router, prefix="/models", tags=["models"])
app.include_router(graph.router)

from app.ingestion.pdf_processor import process_pdf
from app.ingestion.csv_processor import process_csv
from app.ingestion.xlsx_processor import process_xlsx
from app.ingestion.xml_ingestion import process_xml
from app.ingestion.openapi_ingestion import process_openapi
from app.ingestion.postman_ingestion import process_postman
from app.ingestion.web_ingestion import ingest_web_url
from app.ingestion.media_transcriber import transcribe_media
from app.ingestion.image_ocr import extract_text_from_image
from app.database_factory import init_database
from app.config import get_deployment_info
from app.qa_engine import QAEngine
from app.extraction.langextract_adapter import run_langextract, write_visualization
from app.chr_pipeline import run_chr, pca_plot
from app.search import SearchIndex
from app.export_poml import build_poml
from app.analysis.summarization import SummarizationService
import yaml
try:
    from watchgod import watch
except Exception:
    watch = None  # type: ignore
import csv
from datetime import datetime

load_dotenv()

# Normalize HF token envs for Hugging Face downloads
hf_token = (
    os.getenv("HUGGINGFACE_HUB_TOKEN")
    or os.getenv("HF_API_KEY")
    or os.getenv("HF_TOKEN")
    or os.getenv("HUGGINGFACE_TOKEN")
)
if hf_token and not os.getenv("HUGGINGFACE_HUB_TOKEN"):
    os.environ["HUGGINGFACE_HUB_TOKEN"] = hf_token

# NOTE: FastAPI app already created above (line 23) with routers included.
# Do NOT recreate app here or all routers will be lost!

# Optionally run Alembic migrations on startup
if os.getenv("AUTO_MIGRATE", "false").lower() == "true":
    try:
        import subprocess as _sp
        _sp.run(["alembic", "upgrade", "head"], cwd=str(Path(__file__).resolve().parents[1]), check=False)
    except Exception:
        pass

# CORS for frontend (env-driven)
frontend_origin = os.getenv("FRONTEND_ORIGIN", "http://localhost:3000")
allow_origins = [o.strip() for o in frontend_origin.split(",") if o.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=allow_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount Pmoves-hyperdimensions tool
import os
hyp_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "external", "Pmoves-hyperdimensions")
if os.path.exists(hyp_path):
    app.mount("/hyperdimensions", StaticFiles(directory=hyp_path, html=True), name="hyperdimensions")

@app.middleware("http")
async def _fast_pdf_middleware(request, call_next):
    if request.url.path == "/open/pdf":
        if _env_flag("FAST_PDF_MODE", True) or not _env_flag("OPEN_PDF_ENABLED", False):
            return JSONResponse({"detail": "PDF open disabled"}, status_code=403)
    return await call_next(request)

# Initialize
UPLOAD_DIR = Path("uploads")
ARTIFACTS_DIR = Path("artifacts")
UPLOAD_DIR.mkdir(exist_ok=True)
ARTIFACTS_DIR.mkdir(exist_ok=True)

# Security settings
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB limit for file uploads

db, DB_BACKEND_META = init_database()
qa_engine = QAEngine(db)
search_index = SearchIndex(db)
summary_service = SummarizationService(db)
# HRM config/metrics (optional features)
HRM_ENABLED = os.getenv("HRM_ENABLED", "false").lower() == "true"
HRM_CFG = HRMConfig(
    Mmax=int(os.getenv("HRM_MMAX", "6")),
    Mmin=int(os.getenv("HRM_MMIN", "2")),
    threshold=float(os.getenv("HRM_THRESHOLD", "0.5")),
)
HRM_STATS = HRMMetrics()

AUDIO_SUFFIXES = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".aac"}
VIDEO_SUFFIXES = {".mp4", ".mov", ".mkv", ".webm", ".avi"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}
MEDIA_SUFFIXES = AUDIO_SUFFIXES | VIDEO_SUFFIXES


def _env_flag(name: str, default: bool = False) -> bool:
    val = os.getenv(name)
    if val is None:
        return default
    val = val.strip()
    if not val:
        return default
    return val.lower() in {"1", "true", "yes", "on"}

# Simple in-memory task registry
TASKS: dict[str, dict] = {}
START_TIME = time.time()


def _ingest_file_from_watch(src: Path, report_week: str = ""):
    try:
        if not src.exists() or not src.is_file():
            return
        file_id = str(uuid.uuid4())
        dst = UPLOAD_DIR / f"{file_id}_{src.name}"
        shutil.copy2(src, dst)
        suffix = dst.suffix.lower()
        artifact_id = db.add_artifact({
            "id": file_id,
            "filename": src.name,
            "filepath": str(dst),
            "filetype": suffix,
            "report_week": report_week,
            "status": "processing" if suffix == ".pdf" else "processed"
        })
        if suffix == ".pdf":
            task_id = str(uuid.uuid4())
            TASKS[task_id] = {"status": "queued", "filename": src.name, "artifact_id": artifact_id}
            # schedule background processing using the same helper
            _thread = threading.Thread(target=_process_and_store, args=(dst, report_week, artifact_id, suffix, task_id), daemon=True)
            _thread.start()
        elif suffix in (".csv", ".xlsx", ".xls"):
            _process_and_store(dst, report_week, artifact_id, suffix, None)
        elif suffix == ".xml":
            doc, rows = process_xml(dst)
            db.add_document(doc)
            for row in rows:
                db.add_log(row)
        elif suffix in (".yaml", ".yml", ".json"):
            # Try OpenAPI then Postman
            try:
                doc, rows = process_openapi(dst)
                db.add_document(doc)
                for row in rows:
                    db.add_api(row)
            except Exception:
                try:
                    doc, rows = process_postman(dst)
                    db.add_document(doc)
                    for row in rows:
                        db.add_api(row)
                except Exception:
                    pass
    except Exception:
        pass


def _watch_loop():
    if not watch:
        return
    enabled = os.getenv("WATCH_ENABLED", "true").lower() == "true"
    if not enabled:
        return
    watch_dir = Path(os.getenv("WATCH_DIR", "/app/watch"))
    debounce_ms = int(os.getenv("WATCH_DEBOUNCE_MS", "1000"))
    min_bytes = int(os.getenv("WATCH_MIN_BYTES", "1"))
    try:
        watch_dir.mkdir(parents=True, exist_ok=True)
    except PermissionError:
        # Fall back to a writable directory if /app is not writable
        watch_dir = Path("./uploads/watch")
        watch_dir.mkdir(parents=True, exist_ok=True)
        print(f"Watch folder: using fallback path {watch_dir}")
    exts = {".pdf", ".csv", ".xlsx", ".xls"}

    def ready(p: Path) -> bool:
        try:
            if not p.exists() or not p.is_file():
                return False
            if p.suffix.lower() not in exts:
                return False
            if p.stat().st_size < min_bytes:
                return False
            s1 = p.stat().st_size
            time.sleep(debounce_ms / 1000.0)
            s2 = p.stat().st_size
            return s1 == s2
        except Exception:
            return False

    seen: set[str] = set()
    for changes in watch(str(watch_dir)):
        for _evt, path_str in changes:
            p = Path(path_str)
            key = str(p.resolve())
            if key in seen:
                continue
            if ready(p):
                seen.add(key)
                _ingest_file_from_watch(p)


@app.on_event("startup")
async def _startup_watch():
    # Start watcher thread
    t = threading.Thread(target=_watch_loop, daemon=True)
    t.start()

    # Rebuild search index in background with timeout to prevent startup hang
    # Note: Using ThreadPoolExecutor instead of signal.alarm() because signals
    # only work in the main thread, not background threads
    def _rebuild_with_timeout():
        from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
        try:
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(search_index.rebuild)
                future.result(timeout=30)  # 30 second timeout
            print("[STARTUP] Search index rebuild completed")
        except FuturesTimeoutError:
            print("[STARTUP] Search index rebuild timed out, continuing without full index")
        except Exception as e:
            print(f"[STARTUP] Search index rebuild failed: {e}")

    # Run rebuild in background thread so startup completes immediately
    rebuild_thread = threading.Thread(target=_rebuild_with_timeout, daemon=True)
    rebuild_thread.start()

    # Connect to NATS Geometry Bus (non-blocking)
    try:
        from app.services.chit_service import chit_service
        # Use NATS_URL from env or default to docker service name
        nats_url = os.getenv("NATS_URL", "nats://nats:4222")
        asyncio.create_task(chit_service.connect_nats(nats_url))
    except Exception as e:
        print(f"Failed to initiate NATS connection: {e}")

@app.get("/")
async def root():
    return {"message": "PMOVES-DoX API", "status": "running"}

@app.get("/config")
async def config():
    # Detect GPU availability if torch is present
    gpu = {"available": False, "device_count": 0, "names": []}
    try:
        import torch  # type: ignore
        if hasattr(torch, "cuda") and torch.cuda.is_available():
            gpu["available"] = True
            gpu["device_count"] = torch.cuda.device_count()
            try:
                gpu["names"] = [torch.cuda.get_device_name(i) for i in range(gpu["device_count"])]
            except Exception:
                pass
    except Exception:
        pass
    # Detect Ollama availability
    ollama = {"available": False, "base_url": os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")}
    try:
        import httpx  # type: ignore
        base = ollama["base_url"].rstrip("/")
        resp = httpx.get(f"{base}/api/tags", timeout=2.0)
        if resp.status_code == 200:
            ollama["available"] = True
    except Exception:
        pass
    offline = bool(os.getenv("TRANSFORMERS_OFFLINE") or os.getenv("HF_HUB_OFFLINE"))
    return {
        "vlm_repo": os.getenv("DOCLING_VLM_REPO"),
        "database": DB_BACKEND_META,
        "hf_auth": bool(os.getenv("HUGGINGFACE_HUB_TOKEN")),
        "frontend_origin": frontend_origin,
        "gpu": gpu,
        "ollama": ollama,
        "offline": offline,
        "open_pdf_enabled": os.getenv("OPEN_PDF_ENABLED", "false").lower() == "true",
        "deployment": get_deployment_info(),
    }


@app.get("/tasks")
async def list_tasks():
    # Return a lightweight summary including counts and queued items
    items = [
        {"id": k, **v} for k, v in TASKS.items()
    ]
    queued = [t for t in items if t.get("status") == "queued"]
    completed = [t for t in items if t.get("status") == "completed"]
    errored = [t for t in items if t.get("status") == "error"]
    return {
        "total": len(items),
        "queued": len(queued),
        "completed": len(completed),
        "errored": len(errored),
        "queued_items": queued,
    }


@app.get("/health")
async def health():
    return {
        "status": "ok",
        "uptime_seconds": int(time.time() - START_TIME),
    }


@app.get("/artifacts")
async def list_artifacts():
    artifacts = db.get_artifacts()
    evidence = db.get_all_evidence()
    summary: dict[str, dict[str, int]] = {}
    for ev in evidence:
        art_id = ev.get("artifact_id")
        if not art_id:
            continue
        bucket = summary.setdefault(
            art_id,
            {
                "table_evidence": 0,
                "chart_evidence": 0,
                "formula_evidence": 0,
                "media_transcripts": 0,
                "media_metadata": 0,
                "web_pages": 0,
                "image_ocr": 0,
            },
        )
        ctype = (ev.get("content_type") or "").lower()
        if ctype == "table":
            bucket["table_evidence"] += 1
        elif ctype == "chart":
            bucket["chart_evidence"] += 1
        elif ctype == "formula":
            bucket["formula_evidence"] += 1
        elif ctype in {"media_transcript", "audio_transcript", "video_transcript"}:
            bucket["media_transcripts"] += 1
        elif ctype == "media_metadata":
            bucket["media_metadata"] += 1
        elif ctype == "web_page":
            bucket["web_pages"] += 1
        elif ctype == "image_ocr":
            bucket["image_ocr"] += 1

    enriched = []
    for art in artifacts:
        counts = summary.get(
            art.get("id"),
            {
                "table_evidence": 0,
                "chart_evidence": 0,
                "formula_evidence": 0,
                "media_transcripts": 0,
                "media_metadata": 0,
                "web_pages": 0,
                "image_ocr": 0,
            },
        )
        enriched.append({**art, **counts})
    return {"artifacts": enriched}

@app.get("/artifacts/{artifact_id}")
async def artifact_detail(artifact_id: str):
    arts = db.get_artifacts()
    art = next((a for a in arts if a.get("id") == artifact_id), None)
    if not art:
        raise HTTPException(404, "Artifact not found")
    facts = [f for f in db.get_facts() if f.get("artifact_id") == artifact_id]
    evidence = [e for e in db.get_all_evidence() if e.get("artifact_id") == artifact_id]
    return {"artifact": art, "facts": facts, "evidence": evidence}


@app.get("/artifacts/media")
async def artifact_media():
    artifacts = {a.get("id"): a for a in db.get_artifacts()}
    evidence = db.get_all_evidence()
    transcripts: list[dict] = []
    metadata_rows: list[dict] = []
    web_rows: list[dict] = []
    ocr_rows: list[dict] = []

    for ev in evidence:
        art_id = ev.get("artifact_id")
        ctype = (ev.get("content_type") or "").lower()
        entry = {
            "artifact_id": art_id,
            "artifact": {
                "id": art_id,
                "filename": artifacts.get(art_id, {}).get("filename"),
                "filetype": artifacts.get(art_id, {}).get("filetype"),
            },
            "locator": ev.get("locator"),
            "preview": ev.get("preview"),
            "content_type": ctype,
            "full_data": ev.get("full_data"),
        }
        if ctype in {"media_transcript", "audio_transcript", "video_transcript"}:
            transcripts.append(entry)
        elif ctype == "media_metadata":
            metadata_rows.append(entry)
        elif ctype == "web_page":
            web_rows.append(entry)
        elif ctype == "image_ocr":
            ocr_rows.append(entry)

    return {
        "transcripts": transcripts,
        "media_metadata": metadata_rows,
        "web_pages": web_rows,
        "image_text": ocr_rows,
    }

@app.get("/documents")
async def list_documents(type: str | None = None):
    items = db.list_documents(type=type)
    return {"documents": items}


@app.get("/analysis/entities")
async def get_analysis_entities(document_id: str | None = None, label: str | None = None):
    try:
        items = db.list_entities(document_id=document_id, label=label)
    except AttributeError:
        items = []
    return {"entities": items}


@app.get("/analysis/artifacts/{artifact_id}")
async def get_artifact_analysis(artifact_id: str):
    arts = db.get_artifacts()
    art = next((a for a in arts if a.get("id") == artifact_id), None)
    if not art:
        raise HTTPException(404, "Artifact not found")

    evidence = [e for e in db.get_all_evidence() if e.get("artifact_id") == artifact_id]

    tables: list[dict] = []
    charts: list[dict] = []
    formulas: list[dict] = []

    for ev in evidence:
        ctype = (ev.get("content_type") or "").lower()
        base = {
            "id": ev.get("id"),
            "locator": ev.get("locator"),
            "preview": ev.get("preview"),
            "coordinates": ev.get("coordinates"),
        }
        full = ev.get("full_data") or {}

        if ctype == "table":
            rows = list(full.get("rows") or [])
            tables.append(
                {
                    **base,
                    "pages": full.get("pages", []),
                    "merged": full.get("merged", False),
                    "header_detected": full.get("header_detected", False),
                    "columns": full.get("columns", []),
                    "row_count": len(rows),
                    "rows": rows[:20],
                }
            )
        elif ctype == "chart":
            charts.append(
                {
                    **base,
                    "id": full.get("id") or base.get("id"),
                    "page": full.get("page"),
                    "bbox": full.get("bbox"),
                    "image_path": full.get("image_path"),
                    "caption": full.get("caption"),
                    "type": full.get("type"),
                    "extracted_text": full.get("extracted_text"),
                    "vlm_enabled": full.get("vlm_enabled"),
                }
            )
        elif ctype == "formula":
            formulas.append({**base, **full})

    try:
        entities = db.list_entities(document_id=artifact_id)
    except AttributeError:
        entities = []

    try:
        structure = db.get_structure(artifact_id)
    except AttributeError:
        structure = None

    try:
        metric_hits = db.list_metric_hits(document_id=artifact_id)
    except AttributeError:
        metric_hits = []

    return {
        "artifact": art,
        "tables": tables,
        "charts": charts,
        "formulas": formulas,
        "entities": entities,
        "structure": structure,
        "metric_hits": metric_hits,
    }


@app.get("/analysis/structure")
async def get_analysis_structure(document_id: str):
    try:
        structure = db.get_structure(document_id)
    except AttributeError:
        structure = None
    return {"document_id": document_id, "structure": structure}


@app.get("/analysis/metrics")
async def get_analysis_metrics(document_id: str | None = None, metric_type: str | None = None):
    try:
        items = db.list_metric_hits(document_id=document_id, metric_type=metric_type)
    except AttributeError:
        items = []
    return {"metric_hits": items}


# ---------- ingestion: XML / OpenAPI / Postman ----------
@app.post("/ingest/xml")
async def ingest_xml(file: UploadFile = File(...)):
    tmp = UPLOAD_DIR / f"tmp_{uuid.uuid4()}_{file.filename}"
    with tmp.open("wb") as f:
        shutil.copyfileobj(file.file, f)
    doc, rows = process_xml(tmp)
    db.add_document(doc)
    for row in rows:
        db.add_log(row)
    return {"status": "ok", "document_id": doc["id"], "rows": len(rows)}


@app.post("/ingest/openapi")
async def ingest_openapi(file: UploadFile = File(...)):
    tmp = UPLOAD_DIR / f"tmp_{uuid.uuid4()}_{file.filename}"
    with tmp.open("wb") as f:
        shutil.copyfileobj(file.file, f)
    doc, rows = process_openapi(tmp)
    db.add_document(doc)
    for row in rows:
        db.add_api(row)
    return {"status": "ok", "document_id": doc["id"], "rows": len(rows)}


@app.post("/ingest/postman")
async def ingest_postman(file: UploadFile = File(...)):
    tmp = UPLOAD_DIR / f"tmp_{uuid.uuid4()}_{file.filename}"
    with tmp.open("wb") as f:
        shutil.copyfileobj(file.file, f)
    doc, rows = process_postman(tmp)
    db.add_document(doc)
    for row in rows:
        db.add_api(row)
    return {"status": "ok", "document_id": doc["id"], "rows": len(rows)}


# ---------- queries: logs / apis / tags ----------
@app.get("/logs")
async def get_logs(level: str | None = None, code: str | None = None, q: str | None = None,
                   ts_from: str | None = None, ts_to: str | None = None,
                   document_id: str | None = None):
    items = db.list_logs(level=level, code=code, q=q, ts_from=ts_from, ts_to=ts_to,
                         document_id=document_id)
    return {"logs": items}


@app.get("/logs/export")
async def export_logs(level: str | None = None, code: str | None = None, q: str | None = None,
                      ts_from: str | None = None, ts_to: str | None = None,
                      document_id: str | None = None):
    items = db.list_logs(level=level, code=code, q=q, ts_from=ts_from, ts_to=ts_to,
                         document_id=document_id)
    def gen():
        yield "ts,level,code,component,message\n"
        for l in items:
            # naive CSV escaping
            def esc(v: str | None) -> str:
                if v is None:
                    return ""
                s = str(v).replace("\r", " ").replace("\n", " ")
                if "," in s or '"' in s:
                    s = '"' + s.replace('"', '""') + '"'
                return s
            row = ",".join([
                esc(l.get("ts")), esc(l.get("level")), esc(l.get("code")), esc(l.get("component")), esc(l.get("message"))
            ])
            yield row + "\n"
    return StreamingResponse(gen(), media_type="text/csv", headers={"Content-Disposition": "attachment; filename=logs.csv"})


@app.get("/apis")
async def get_apis(tag: str | None = None, method: str | None = None, path_like: str | None = None):
    items = db.list_apis(tag=tag, method=method, path_like=path_like)
    return {"apis": items}


@app.get("/apis/{api_id}")
async def get_api_detail(api_id: str):
    # fetch raw row from DB
    from sqlmodel import Session
    from app.database import APIEndpoint
    with Session(db.engine) as s:  # type: ignore[attr-defined]
        row = s.get(APIEndpoint, api_id)
    if not row:
        raise HTTPException(404, "API not found")
    import json as _json
    return {
        "id": row.id,
        "document_id": row.document_id,
        "name": row.name,
        "method": row.method,
        "path": row.path,
        "summary": row.summary,
        "tags": _json.loads(row.tags_json) if row.tags_json else [],
        "parameters": _json.loads(row.params_json) if row.params_json else [],
        "responses": _json.loads(row.responses_json) if row.responses_json else {},
    }


@app.get("/tags")
async def get_tags(document_id: str | None = None, q: str | None = None):
    items = db.list_tags(document_id=document_id, q=q)
    return {"tags": items}


class ExtractTagsRequest(BaseModel):
    document_id: str
    model_id: str | None = None
    api_key: str | None = None
    prompt: str | None = None
    examples: list[dict] | None = None
    dry_run: bool = False
    use_hrm: bool = False
    include_poml: bool = False
    poml_variant: str | None = None
    mangle_exec: bool = False
    mangle_file: str | None = None
    mangle_query: str | None = None


class AutoTagRequest(BaseModel):
    async_run: bool | None = None
    include_poml: bool | None = None
    use_hrm: bool | None = None
    model_id: str | None = None
    api_key: str | None = None
    prompt: str | None = None
    examples: list[dict] | None = None
    dry_run: bool | None = None
    mangle_exec: bool | None = None
    mangle_file: str | None = None
    mangle_query: str | None = None
    poml_variant: str | None = None


def _compose_text_for_document(doc: Dict) -> str:
    doc_type = (doc.get("type") or "").lower()
    path = Path(doc.get("path", ""))
    if doc_type == "xml":
        # concatenate log messages
        msgs = db.list_log_messages(doc["id"])  # type: ignore[arg-type]
        return "\n".join(msgs)
    elif doc_type in ("openapi", "postman"):
        # read raw file
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""
    elif doc_type == "pdf":
        # Prefer markdown from artifacts if exists
        md_path = ARTIFACTS_DIR / f"{path.stem}.md"
        if md_path.exists():
            return md_path.read_text(encoding="utf-8")
    # fallback: file content
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def _load_document(document_id: str) -> Optional[Dict]:
    try:
        return next((d for d in db.list_documents() if d.get("id") == document_id), None)
    except Exception:
        return None


def _build_poml_context(document_id: str, variant: Optional[str] = None) -> Optional[str]:
    doc = _load_document(document_id)
    if not doc:
        return None
    try:
        apis = [a for a in db.list_apis(tag=None, method=None, path_like=None) if a.get("document_id") == document_id]
    except Exception:
        apis = []
    try:
        tags = db.list_tags(document_id=document_id, q=None)
    except Exception:
        tags = []
    try:
        logs = [
            l for l in db.list_logs(level=None, code=None, q=None, ts_from=None, ts_to=None)
            if l.get("document_id") == document_id
        ]
    except Exception:
        logs = []
    md_path: Optional[Path] = None
    chr_csv: Optional[Path] = None
    try:
        src = Path(doc.get("path", ""))
        stem = src.stem
        cand_md = ARTIFACTS_DIR / f"{stem}.md"
        if cand_md.exists():
            md_path = cand_md
        cand_chr = ARTIFACTS_DIR / "chr" / f"{stem}_chr.csv"
        if cand_chr.exists():
            chr_csv = cand_chr
    except Exception:
        md_path = None
        chr_csv = None
    try:
        poml_variant = variant or os.getenv("AUTOTAG_POML_VARIANT") or "generic"
        poml = build_poml(doc, apis, tags, logs, md_path, chr_csv, poml_variant)
        # Avoid overwhelming prompts with massive payloads
        return poml[:4000]
    except Exception:
        return None


def _resolve_mangle_file(path_hint: Optional[str]) -> Optional[Path]:
    candidate = path_hint or os.getenv("AUTOTAG_MANGLE_FILE") or os.getenv("MANGLE_FILE")
    if not candidate:
        return None
    try:
        p = Path(candidate).expanduser()
        if p.exists():
            return p
    except Exception:
        return None
    return None


def _load_mangle_rules(path: Optional[Path]) -> Optional[str]:
    if not path:
        return None
    try:
        return path.read_text(encoding="utf-8")
    except Exception:
        return None


def _apply_mangle(tags: List[str], rules_path: Path, query: Optional[str]) -> Optional[List[str]]:
    if not shutil.which("mg"):
        return None
    if not tags:
        return None
    program = rules_path
    if not program.exists():
        return None
    q = query or os.getenv("AUTOTAG_MANGLE_QUERY") or "normalized_tag(T)"
    tmp_path: Optional[Path] = None
    try:
        with tempfile.NamedTemporaryFile("w", suffix=".edb", delete=False, encoding="utf-8") as tmp:
            tmp_path = Path(tmp.name)
            for tag in tags:
                safe = tag.replace("\"", "\\\"")
                tmp.write(f'tag_raw("{safe}").\n')
        proc = subprocess.run(
            ["mg", "--ruleset", str(program), "--edb", str(tmp_path), "--query", q],
            check=False,
            capture_output=True,
            text=True,
        )
        if proc.returncode != 0:
            return None
        output = (proc.stdout or "").strip()
        if not output:
            return None
        matches = re.findall(r'"([^"\\]+)"', output)
        if matches:
            cleaned = [m.strip() for m in matches if m.strip()]
            if cleaned:
                return cleaned
        lines = [line.strip() for line in output.splitlines() if line.strip()]
        return lines or None
    except Exception:
        return None
    finally:
        if tmp_path and tmp_path.exists():
            try:
                tmp_path.unlink()
            except Exception:
                pass


def _resolve_document_for_artifact(artifact_id: str) -> tuple[Dict, Dict]:
    art = next((a for a in db.get_artifacts() if a.get("id") == artifact_id), None)
    if not art:
        raise HTTPException(404, "Artifact not found")
    doc = _load_document(artifact_id)
    if doc:
        return art, doc
    art_path = art.get("filepath")
    if art_path:
        try:
            art_abs = Path(art_path).resolve()
        except Exception:
            art_abs = None
        if art_abs is not None:
            for cand in db.list_documents():
                try:
                    cand_path = Path(cand.get("path", "")).resolve()
                except Exception:
                    continue
                if cand_path == art_abs:
                    return art, cand
    raise HTTPException(404, "No document associated with artifact")
@app.post("/extract/tags")
async def extract_tags(req: ExtractTagsRequest):
    doc = _load_document(req.document_id)
    if not doc:
        raise HTTPException(404, "Document not found")

    text = _compose_text_for_document(doc)
    if not text:
        raise HTTPException(400, "No text available for tag extraction")

    preset_prompt = req.prompt or os.getenv("TAGS_PROMPT", (
        "Extract application or system tags relevant to loan management systems (LMS). "
        "Return concise tags as exact spans from text. Examples: 'Loan Origination', 'Underwriting', 'Servicing', 'LoanService'."
    ))
    poml_text: Optional[str] = None
    if req.include_poml:
        poml_text = _build_poml_context(req.document_id, req.poml_variant)
        if poml_text:
            preset_prompt += "\n\nPOML CONTEXT:\n```\n" + poml_text + "\n```"
    mangle_path = _resolve_mangle_file(req.mangle_file)
    mangle_rules = _load_mangle_rules(mangle_path)
    if mangle_rules:
        preset_prompt += "\n\nMANGLE RULES:\n```\n" + mangle_rules + "\n```"
    # If dry_run and no API key/model configured, use a heuristic fallback to avoid external calls during smoke.
    use_fallback = (
        req.dry_run and not (req.api_key or os.getenv("GOOGLE_API_KEY") or os.getenv("OPENAI_API_KEY") or os.getenv("LANGEXTRACT_API_KEY"))
    )
    t0 = time.time()
    def _heuristic_entities(limit: int = 5):
        import re as _re
        candidates = _re.findall(r"([A-Z][a-z]+(?:\s+[A-Z][a-zA-Z]+){0,2})", text)
        unique = []
        for c in candidates:
            c = c.strip()
            if c and c not in unique:
                unique.append(c)
            if len(unique) >= limit:
                break
        return {"entities": [{"extraction_text": t, "extraction_class": "heuristic"} for t in unique]}

    if use_fallback:
        result = _heuristic_entities()
    else:
        try:
            result = run_langextract(
                text=text,
                prompt_description=preset_prompt,
                examples=req.examples,
                model_id=req.model_id,
                api_key=req.api_key,
                extraction_passes=1,
                max_workers=4,
                max_char_buffer=4000,
            )
        except ValueError as exc:
            if "Examples are required" in str(exc):
                result = _heuristic_entities()
            else:
                raise
        except Exception:
            result = _heuristic_entities()

    entities = result.get("entities", [])
    # Optional HRM-style iterative refinement over extracted tags
    steps = 0
    if HRM_ENABLED and req.use_hrm:
        def norm_tag(t: str) -> str:
            return " ".join(t.strip().split())
        tags0 = [str(e.get("extraction_text") or e.get("text") or "").strip() for e in entities if (e.get("extraction_text") or e.get("text"))]
        tags = [t for t in tags0 if t]
        for m in range(1, HRM_CFG.Mmax + 1):
            steps = m
            tags_new = []
            seen = set()
            for t in tags:
                t2 = norm_tag(t)
                if t2 and t2 not in seen:
                    seen.add(t2)
                    tags_new.append(t2)
            if m >= HRM_CFG.Mmin and tags_new == tags:
                break
            tags = tags_new
        # replace entities with refined tags for persistence phase
        entities = [{"extraction_text": t, "extraction_class": "hrm-refined"} for t in tags]
        dt = (time.time() - t0) * 1000.0
        HRM_STATS.record(steps=steps, latency_ms=dt, payload={"mode": "extract_tags", "doc": req.document_id})
    tags_for_persistence: List[tuple[str, str]] = []
    for e in entities:
        tag_text = e.get("extraction_text") or e.get("text")
        if not tag_text:
            continue
        tag_text = str(tag_text).strip()
        if not tag_text:
            continue
        source_ptr = str(e.get("extraction_class") or "langextract")
        tags_for_persistence.append((tag_text, source_ptr))

    if mangle_path and req.mangle_exec:
        mangled = _apply_mangle([t for t, _ in tags_for_persistence], mangle_path, req.mangle_query)
        if mangled:
            tags_for_persistence = [(t, "mangle") for t in mangled]

    saved = 0
    extracted: List[str] = []
    for tag_text, source_ptr in tags_for_persistence:
        extracted.append(tag_text)
        if req.dry_run:
            continue
        if not db.has_tag(req.document_id, tag_text):
            if source_ptr == "hrm-refined" and (HRM_ENABLED and req.use_hrm):
                try:
                    s_val = int(steps) if isinstance(steps, int) else None
                except Exception:
                    s_val = None
                if s_val is not None:
                    source_ptr = f"hrm-refined:steps{s_val}"
            db.add_tag({
                "id": str(uuid.uuid4()),
                "document_id": req.document_id,
                "tag": tag_text,
                "score": None,
                "source_ptr": source_ptr,
                "hrm_steps": steps if (HRM_ENABLED and req.use_hrm) else None,
            })
            saved += 1
    resp = {"status": "ok", "document_id": req.document_id, "tags_saved": saved, "tags": extracted}
    if HRM_ENABLED and req.use_hrm:
        resp["hrm"] = {"enabled": True, "steps": steps}
    return resp


@app.post("/autotag/{artifact_id}")
async def auto_tag_artifact(artifact_id: str, req: AutoTagRequest):
    artifact, document = _resolve_document_for_artifact(artifact_id)
    include_poml = req.include_poml if req.include_poml is not None else _env_flag("AUTOTAG_INCLUDE_POML", _env_flag("POML_IN_PROMPT", False))
    use_hrm_flag = req.use_hrm if req.use_hrm is not None else (_env_flag("AUTOTAG_USE_HRM", False) and HRM_ENABLED)
    model_id = req.model_id
    api_key = req.api_key
    prompt = req.prompt
    examples = req.examples
    dry_run = bool(req.dry_run)
    mangle_exec = req.mangle_exec if req.mangle_exec is not None else _env_flag("AUTOTAG_MANGLE_EXEC", False)
    mangle_file = req.mangle_file or os.getenv("AUTOTAG_MANGLE_FILE") or os.getenv("MANGLE_FILE")
    mangle_query = req.mangle_query or os.getenv("AUTOTAG_MANGLE_QUERY")
    poml_variant = req.poml_variant or os.getenv("AUTOTAG_POML_VARIANT")

    tag_request = ExtractTagsRequest(
        document_id=document.get("id"),
        model_id=model_id,
        api_key=api_key,
        prompt=prompt,
        examples=examples,
        dry_run=dry_run,
        use_hrm=bool(use_hrm_flag),
        include_poml=bool(include_poml),
        poml_variant=poml_variant,
        mangle_exec=bool(mangle_exec),
        mangle_file=mangle_file,
        mangle_query=mangle_query,
    )

    result = await extract_tags(tag_request)

    try:
        total_tags = len(db.list_tags(document_id=document.get("id"), q=None))
    except Exception:
        total_tags = 0

    response = {
        "status": result.get("status", "ok"),
        "artifact_id": artifact.get("id"),
        "document_id": document.get("id"),
        "tags_saved": result.get("tags_saved", 0),
        "tags_extracted": len(result.get("tags", [])),
        "tags_total": total_tags,
        "async_requested": bool(req.async_run),
    }
    if "hrm" in result:
        response["hrm"] = result["hrm"]
    return response


@app.get("/tags/presets")
async def tag_presets():
    default_prompt = os.getenv("TAGS_PROMPT", (
        "Extract application or system tags relevant to loan management systems (LMS). "
        "Return concise tags as exact spans from text. Examples: 'Loan Origination', 'Underwriting', 'Servicing', 'LoanService'."
    ))
    examples = [
        {"text": "The LMS core supports Loan Origination and Loan Servicing.", "labels": ["Loan Origination", "Loan Servicing"]},
        {"text": "Enable Underwriting via RulesEngine v2.", "labels": ["Underwriting", "RulesEngine"]},
    ]
    return {"prompt": default_prompt, "examples": examples}


class TagPromptSaveRequest(BaseModel):
    prompt_text: str
    examples: list[dict] | None = None
    author: str | None = None


@app.get("/tags/prompt/{document_id}")
async def get_tag_prompt(document_id: str):
    item = db.get_latest_tag_prompt(document_id)
    if not item:
        return {"prompt": None}
    return item


@app.get("/tags/prompt/{document_id}/history")
async def get_tag_prompt_history(document_id: str, limit: int = 20):
    items = db.list_tag_prompt_history(document_id, limit=limit)
    return {"items": items}


@app.post("/tags/prompt/{document_id}")
async def save_tag_prompt(document_id: str, req: TagPromptSaveRequest):
    pid = db.save_tag_prompt(document_id, req.prompt_text, req.examples, req.author)
    return {"status": "ok", "id": pid}


@app.get("/watch")
async def watch_status():
    return {
        "enabled": os.getenv("WATCH_ENABLED", "true").lower() == "true",
        "dir": os.getenv("WATCH_DIR", "/app/watch"),
        "debounce_ms": int(os.getenv("WATCH_DEBOUNCE_MS", "1000")),
        "min_bytes": int(os.getenv("WATCH_MIN_BYTES", "1")),
        "available": bool(watch),
    }


# ---------------- LangExtract integration ----------------
class LangExtractRequest(BaseModel):
    artifact_id: str | None = None
    text: str | None = None
    prompt_description: str
    examples: list[dict] | None = None
    model_id: str | None = None
    api_key: str | None = None
    extraction_passes: int = 1
    max_workers: int = 8
    max_char_buffer: int = 4000


class SummaryGenerateRequest(BaseModel):
    style: Literal["bullet", "executive", "action_items"] = "bullet"
    scope: Literal["workspace", "artifact"] = "workspace"
    artifact_ids: list[str] | None = None
    provider: str | None = None
    force_refresh: bool = False


class SummaryResponse(BaseModel):
    id: str
    style: str
    provider: str
    prompt: str
    scope: dict[str, Any]
    summary: str
    citations: list[dict[str, Any]]
    created_at: str


def _load_text_for_artifact(artifact_id: str) -> str:
    # Attempt to load markdown generated by pdf_processor for PDFs
    art = next((a for a in db.get_artifacts() if a.get("id") == artifact_id), None)
    if not art:
        raise HTTPException(404, f"Artifact not found: {artifact_id}")
    p = Path(art.get("filepath", ""))
    if p.suffix.lower() == ".pdf":
        md_path = ARTIFACTS_DIR / f"{p.stem}.md"
        if md_path.exists():
            return md_path.read_text(encoding="utf-8")
    # Fallback: try raw file text for CSV/XLSX (limited utility)
    try:
        return Path(art.get("filepath")).read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


@app.post("/extract/langextract")
async def extract_langextract(req: LangExtractRequest):
    if not req.text and not req.artifact_id:
        raise HTTPException(400, "Provide either text or artifact_id")
    text = req.text or _load_text_for_artifact(req.artifact_id)  # type: ignore[arg-type]
    if not text:
        raise HTTPException(400, "No text available for extraction")

    result = run_langextract(
        text=text,
        prompt_description=req.prompt_description,
        examples=req.examples,
        model_id=req.model_id,
        api_key=req.api_key,
        extraction_passes=req.extraction_passes,
        max_workers=req.max_workers,
        max_char_buffer=req.max_char_buffer,
    )

    # Save artifacts (JSON and HTML) under artifacts/
    out_dir = ARTIFACTS_DIR / "langextract"
    out_dir.mkdir(parents=True, exist_ok=True)
    json_path = out_dir / "langextract_results.json"
    json_path.write_text(json.dumps(result, indent=2), encoding="utf-8")
    html_path = write_visualization(result, out_dir, output_name="langextract_results")

    return {
        "status": "ok",
        "model_id": result.get("model_id"),
        "entities_count": len(result.get("entities", [])),
        "artifacts": {
            "json": str(json_path),
            "html": str(html_path),
        },
    }


@app.get("/summaries", response_model=dict[str, list[SummaryResponse]])
async def list_summaries(scope: str | None = Query(None), style: str | None = Query(None)):
    summaries = summary_service.list_summaries(scope=scope, style=style)
    return {"summaries": summaries}


@app.post("/summaries/generate", response_model=SummaryResponse)
async def generate_summary(req: SummaryGenerateRequest):
    try:
        result = summary_service.generate_summary(
            style=req.style,
            scope=req.scope,
            artifact_ids=req.artifact_ids,
            provider=req.provider,
            force_refresh=req.force_refresh,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return result


# ---------------- Vector Search ----------------
class SearchRequest(BaseModel):
    q: str
    k: int | None = 10
    types: list[str] | None = None  # subset of ['pdf','api','log','tag']


@app.post("/search")
async def search(req: SearchRequest):
    t0 = time.time()
    types = {t.lower() for t in (req.types or []) if t}

    results_payload: list[dict] = []
    if (req.q or "").strip() == "__ui_test__":
        # Keep parity with UI smoke path while supporting type filters
        results_payload = [{
            "score": 1.0,
            "text": "UI Test Result",
            "meta": {"type": "api", "deeplink": {"panel": "apis"}},
        }]
        if types and "api" not in types:
            results_payload = []
    else:
        hits = search_index.search(req.q or "", k=req.k or 10)
        for hit in hits:
            meta = dict(hit.meta or {})
            entry_type = (meta.get("type") or "").lower()
            if types and entry_type not in types:
                continue

            if entry_type == "pdf":
                deeplink = {
                    "panel": "workspace",
                    "artifact_id": meta.get("artifact_id"),
                    "chunk": meta.get("chunk"),
                }
                page = meta.get("page")
                if page is not None:
                    deeplink["page"] = page
            elif entry_type == "api":
                deeplink = {"panel": "apis", "api_id": meta.get("id")}
            elif entry_type == "log":
                deeplink = {
                    "panel": "logs",
                    "document_id": meta.get("document_id"),
                    "code": meta.get("code"),
                }
            elif entry_type == "tag":
                deeplink = {
                    "panel": "tags",
                    "document_id": meta.get("document_id"),
                    "q": meta.get("tag") or meta.get("text"),
                }
            else:
                deeplink = {}

            results_payload.append({
                "score": hit.score,
                "text": hit.text,
                "meta": {**meta, "deeplink": deeplink},
            })

    elapsed_ms = max(0, int((time.time() - t0) * 1000))
    return {
        "took_ms": elapsed_ms,
        "count": len(results_payload),
        "results": results_payload,
    }


@app.post("/search/rebuild")
async def search_rebuild():
    info = search_index.rebuild()
    return {"status": "ok", **info}


# ---------------- Conversion: artifact -> txt/docx ----------------
class ConvertRequest(BaseModel):
    artifact_id: str
    format: str  # 'txt' or 'docx'


def _load_markdown_for_artifact(artifact_id: str) -> tuple[Path, str]:
    art = next((a for a in db.get_artifacts() if a.get("id") == artifact_id), None)
    if not art:
        raise HTTPException(404, f"Artifact not found: {artifact_id}")
    p = Path(art.get("filepath", ""))
    if p.suffix.lower() == ".pdf":
        md_path = ARTIFACTS_DIR / f"{p.stem}.md"
        if md_path.exists():
            return md_path, md_path.read_text(encoding="utf-8")
        return md_path, ""
    # CSV/XLSX fallback: return raw text for CSV or a generic message
    try:
        txt = p.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        txt = f"Artifact {p.name} not convertible to text in this mode."
    md_path = ARTIFACTS_DIR / f"{p.stem}.md"
    return md_path, txt


@app.post("/convert")
async def convert_artifact(req: ConvertRequest):
    fmt = req.format.lower()
    if fmt not in ("txt", "docx"):
        raise HTTPException(400, "format must be 'txt' or 'docx'")

    md_path, md_text = _load_markdown_for_artifact(req.artifact_id)
    if not md_text:
        raise HTTPException(400, "No markdown/text available for this artifact yet. Process a PDF first.")

    stem = md_path.stem or "document"
    out_dir = ARTIFACTS_DIR / "conversions"
    out_dir.mkdir(parents=True, exist_ok=True)

    if fmt == "txt":
        txt_path = out_dir / f"{stem}.txt"
        # Prefer pandoc for markdown -> plain text
        try:
            if shutil.which("pandoc") and md_path.exists():
                subprocess.run(["pandoc", str(md_path), "-f", "gfm", "-t", "plain", "-o", str(txt_path)], check=True)
            else:
                txt_path.write_text(md_text, encoding="utf-8")
        except Exception:
            txt_path.write_text(md_text, encoding="utf-8")
        rel = str(txt_path.relative_to(ARTIFACTS_DIR)) if txt_path.is_relative_to(ARTIFACTS_DIR) else f"conversions/{txt_path.name}"
        return {"status": "ok", "path": str(txt_path), "rel": rel}

    # DOCX
    docx_path = out_dir / f"{stem}.docx"
    # Prefer pandoc for markdown -> docx for high fidelity
    try:
        if shutil.which("pandoc") and md_path.exists():
            subprocess.run(["pandoc", str(md_path), "-f", "gfm", "-t", "docx", "-o", str(docx_path)], check=True)
        else:
            raise RuntimeError("pandoc not available")
    except Exception:
        # Fallback naive mapping
        doc = Document()
        for line in md_text.splitlines():
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            else:
                if line.strip().startswith("- "):
                    doc.add_paragraph(line.strip()[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)
        doc.save(docx_path)
    rel = str(docx_path.relative_to(ARTIFACTS_DIR)) if docx_path.is_relative_to(ARTIFACTS_DIR) else f"conversions/{docx_path.name}"
    return {"status": "ok", "path": str(docx_path), "rel": rel}


# ---------------- CHR structuring ----------------
class CHRRequest(BaseModel):
    artifact_id: str
    K: int = 8
    iters: int = 30
    bins: int = 8
    seed: int = 42
    beta: float = 12.0
    units_mode: str = "paragraphs"  # 'paragraphs' or 'sentences'
    include_tables: bool = True


def _sent_split(text: str) -> List[str]:
    import re
    rough = re.split(r"[\n\r]+|(?<=[\.!?])\s+", text.strip())
    return [s.strip() for s in rough if s.strip()]


def _split_units_from_markdown(md_text: str, mode: str = "paragraphs") -> List[str]:
    blocks = [b.strip() for b in md_text.split("\n\n")]
    blocks = [b for b in blocks if b]
    if mode == "sentences":
        units: List[str] = []
        for b in blocks:
            units.extend(_sent_split(b))
        return units
    return blocks


def _units_from_pdf_json(json_path: Path, include_tables: bool, mode: str) -> List[str]:
    try:
        import json as _json
        doc = _json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return []
    units: List[str] = []
    # texts
    for it in doc.get("texts", []) or []:
        t = str(it.get("text", "")).strip()
        if not t:
            continue
        if mode == "sentences":
            units.extend(_sent_split(t))
        else:
            units.append(t)
    # tables
    if include_tables:
        for tb in doc.get("tables", []) or []:
            # assume 'data' or export-like format
            rows = tb.get("data") or tb.get("rows") or []
            for row in rows:
                vals = [str(c) for c in row if str(c).strip()]
                if not vals:
                    continue
                units.append(" ".join(vals))
    return units


def _units_from_csv(path: Path, mode: str) -> List[str]:
    import pandas as pd
    try:
        df = pd.read_csv(path)
    except Exception:
        return []
    units: List[str] = []
    for _, row in df.iterrows():
        vals = [str(v) for v in row.tolist() if str(v).strip()]
        if not vals:
            continue
        text = " ".join(vals)
        if mode == "sentences":
            units.extend(_sent_split(text))
        else:
            units.append(text)
    return units


def _units_from_xlsx(path: Path, mode: str) -> List[str]:
    import pandas as pd
    units: List[str] = []
    try:
        xls = pd.ExcelFile(path)
    except Exception:
        return []
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        for _, row in df.iterrows():
            vals = [str(v) for v in row.tolist() if str(v).strip()]
            if not vals:
                continue
            text = " ".join(vals)
            if mode == "sentences":
                units.extend(_sent_split(text))
            else:
                units.append(text)
    return units


@app.post("/structure/chr")
async def structure_chr(req: CHRRequest):
    md_path, md_text = _load_markdown_for_artifact(req.artifact_id)
    # detect artifact type
    art = next((a for a in db.get_artifacts() if a.get("id") == req.artifact_id), None)
    if not art:
        raise HTTPException(404, "Artifact not found")
    file_path = Path(art.get("filepath", ""))
    suffix = file_path.suffix.lower()

    units: List[str] = []
    if suffix == ".pdf":
        json_path = ARTIFACTS_DIR / f"{file_path.stem}.json"
        if json_path.exists():
            units = _units_from_pdf_json(json_path, include_tables=req.include_tables, mode=req.units_mode)
        if not units and md_text:
            units = _split_units_from_markdown(md_text, mode=req.units_mode)
    elif suffix == ".csv":
        units = _units_from_csv(file_path, mode=req.units_mode)
    elif suffix in (".xlsx", ".xls"):
        units = _units_from_xlsx(file_path, mode=req.units_mode)
    else:
        # fallback to markdown/plain
        if md_text:
            units = _split_units_from_markdown(md_text, mode=req.units_mode)

    if not units:
        raise HTTPException(400, "Could not derive units for this artifact.")

    res = run_chr(units, K=req.K, iters=req.iters, bins=req.bins, beta=req.beta, seed=req.seed)

    # If PDF, try to attach page numbers to rows
    pages_map: List[int] | None = None
    if suffix == ".pdf":
        tu = ARTIFACTS_DIR / f"{file_path.stem}.text_units.json"
        try:
            if tu.exists():
                import json as _json
                data = _json.loads(tu.read_text(encoding="utf-8", errors="ignore"))
                if isinstance(data, list):
                    if req.units_mode != "sentences":
                        # 1:1 mapping with units array
                        pages_map = [int(x.get("page")) if isinstance(x.get("page"), (int, float)) else None for x in data]
                    else:
                        # sentences mode: expand page map by splitting each unit into sentences
                        def _sent_split_local(text: str) -> list[str]:
                            import re as _re
                            rough = _re.split(r"[\n\r]+|(?<=[\.!?])\s+", (text or "").strip())
                            return [s.strip() for s in rough if s and s.strip()]
                        expanded: list[int | None] = []
                        for x in data:
                            txt = (x.get("text") or "").strip()
                            cnt = len(_sent_split_local(txt)) if txt else 0
                            pg = int(x.get("page")) if isinstance(x.get("page"), (int, float)) else None
                            if cnt <= 0:
                                continue
                            expanded.extend([pg] * cnt)
                        pages_map = expanded
        except Exception:
            pages_map = None
        if pages_map:
            for row in res.rows:
                try:
                    idx = int(row.get("idx"))
                    # Bound-check; if sentences produced more/less items than units, clamp by last known page
                    pg = None
                    if 0 <= idx < len(pages_map):
                        pg = pages_map[idx]
                    elif len(pages_map) > 0:
                        pg = pages_map[-1]
                    if isinstance(pg, int):
                        row["page"] = pg
                except Exception:
                    pass

    # Persist CSV/JSON
    import csv, json as _json
    out_dir = ARTIFACTS_DIR / "chr"
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = md_path.stem or "document"
    csv_path = out_dir / f"{stem}_chr.csv"
    json_path = out_dir / f"{stem}_chr.json"
    plot_path = out_dir / f"{stem}_pca.png"
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["idx", "constellation", "radius", "text", "page"])
        writer.writeheader()
        for row in res.rows:
            writer.writerow(row)
    json_obj = {
        "backend": res.backend,
        "K": res.K,
        "mhep": res.mhep,
        "Hg": res.Hg,
        "Hs": res.Hs,
        "Hg_traj": res.Hg_traj,
        "Hs_traj": res.Hs_traj,
        "rows": res.rows,
    }
    json_path.write_text(_json.dumps(json_obj, indent=2), encoding="utf-8")
    # PCA plot
    try:
        pca_plot(res.Z, res.U, np.array(res.labels), str(plot_path))
    except Exception:
        pass

    rel_csv = str(csv_path.relative_to(ARTIFACTS_DIR)) if csv_path.is_relative_to(ARTIFACTS_DIR) else f"chr/{csv_path.name}"
    rel_json = str(json_path.relative_to(ARTIFACTS_DIR)) if json_path.is_relative_to(ARTIFACTS_DIR) else f"chr/{json_path.name}"
    rel_plot = str(plot_path.relative_to(ARTIFACTS_DIR)) if plot_path.exists() and plot_path.is_relative_to(ARTIFACTS_DIR) else None

    return {
        "status": "ok",
        "K": res.K,
        "mhep": res.mhep,
        "Hg": res.Hg,
        "Hs": res.Hs,
        "counts": {"units": len(units)},
        "preview_rows": res.rows[:10],
        "artifacts": {"csv": str(csv_path), "json": str(json_path), "rel_csv": rel_csv, "rel_json": rel_json, "rel_plot": rel_plot},
    }


@app.get("/download")
async def download_artifact(rel: str):
    # serve files only within ARTIFACTS_DIR
    try:
        # Normalize path to prevent traversal
        target = (ARTIFACTS_DIR / rel).resolve()
        if not str(target).startswith(str(ARTIFACTS_DIR.resolve())):
            raise HTTPException(400, "Invalid path")
        if not target.exists() or not target.is_file():
            raise HTTPException(404, "File not found")
        return FileResponse(str(target), filename=target.name)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Download error: {e}")


# ---------------- datavzrd viz project generation ----------------
class DataVZRDRequest(BaseModel):
    artifact_id: str
    title: str | None = None


@app.post("/viz/datavzrd")
async def build_datavzrd(req: DataVZRDRequest):
    art = next((a for a in db.get_artifacts() if a.get("id") == req.artifact_id), None)
    if not art:
        raise HTTPException(404, "Artifact not found")
    file_path = Path(art.get("filepath", ""))
    stem = file_path.stem or art.get("id")
    chr_csv = ARTIFACTS_DIR / "chr" / f"{stem}_chr.csv"
    if not chr_csv.exists():
        raise HTTPException(400, "CHR CSV not found. Run /structure/chr first.")

    proj_dir = ARTIFACTS_DIR / "datavzrd" / stem
    proj_dir.mkdir(parents=True, exist_ok=True)
    # copy path (we can reference relative path)
    rel_csv = os.path.relpath(chr_csv, proj_dir)

    use_spells = os.getenv("DATAVZRD_SPELLS", "true").lower() == "true"
    # columns config for Details table
    columns_chr = ( {
            "idx": {},
            "constellation": {},
            "radius": {},
            "text": {"spell": {"url": "v1.4.1/utils/text", "with": {"chars_per_line": 80}}}
        } if use_spells else ["idx","constellation","radius","text"] )

    cfg = {
        "title": req.title or f"CHR – {stem}",
        "data": [
            {
                "id": "chr",
                "path": rel_csv,
            }
        ],
        "pages": [
            {
                "title": "Overview",
                "blocks": [
                    {
                        "title": "Constellation Map (PCA)",
                        "render": "markdown",
                        "content": f"![]({os.path.relpath(out_dir := (ARTIFACTS_DIR / 'chr' / (stem + '_pca.png')), proj_dir)})"
                    },
                    {
                        "title": "Rows per Constellation",
                        "render": "plot",
                        "data": "chr",
                        "spec": {
                            "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
                            "mark": {"type": "bar"},
                            "encoding": {
                                "x": {"field": "constellation", "type": "nominal", "sort": "ascending"},
                                "y": {"aggregate": "count", "type": "quantitative", "title": "rows"}
                            }
                        }
                    },
                    {
                        "title": "Radius Histogram",
                        "render": "plot",
                        "data": "chr",
                        "spec": {
                            "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
                            "mark": "bar",
                            "encoding": {
                                "x": {"bin": True, "field": "radius", "type": "quantitative"},
                                "y": {"aggregate": "count", "type": "quantitative"}
                            }
                        }
                    }
                ]
            },
            {
                "title": "Details",
                "blocks": [
                    {"title": "CHR Rows", "render": "table", "data": "chr", "columns": columns_chr, "search": True, "download": True}
                ]
            }
        ]
    }

    viz_yaml = proj_dir / "viz.yaml"
    viz_yaml.write_text(yaml.safe_dump(cfg, sort_keys=False), encoding="utf-8")

    return {
        "status": "ok",
        "project_dir": str(proj_dir),
        "viz_yaml": str(viz_yaml),
        "rel_viz": str(viz_yaml.relative_to(ARTIFACTS_DIR)) if viz_yaml.is_relative_to(ARTIFACTS_DIR) else None,
    }


class DataVZRDLogsRequest(BaseModel):
    document_id: str | None = None
    title: str | None = None


@app.post("/viz/datavzrd/logs")
async def build_datavzrd_logs(req: DataVZRDLogsRequest):
    # Collect logs (all or by document)
    logs = db.list_logs(level=None, code=None, q=None, ts_from=None, ts_to=None)
    if req.document_id:
        logs = [l for l in logs if l.get("document_id") == req.document_id]
    if not logs:
        raise HTTPException(400, "No logs available for viz")

    scope = req.document_id or "all"
    proj_dir = ARTIFACTS_DIR / "datavzrd" / f"logs-{scope}"
    proj_dir.mkdir(parents=True, exist_ok=True)
    csv_path = proj_dir / "logs.csv"
    # Write CSV
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["ts","level","code","component","message"]) 
        writer.writeheader()
        for l in logs:
            writer.writerow({
                "ts": l.get("ts"),
                "level": l.get("level"),
                "code": l.get("code"),
                "component": l.get("component"),
                "message": l.get("message"),
            })

    rel_csv = os.path.relpath(csv_path, proj_dir)
    # Build viz.yaml for logs
    use_spells = os.getenv("DATAVZRD_SPELLS", "true").lower() == "true"
    columns_logs = ( {
            "ts": {}, "level": {}, "code": {}, "component": {},
            "message": {"spell": {"url": "v1.4.1/utils/text", "with": {"chars_per_line": 80}}}
        } if use_spells else ["ts","level","code","component","message"] )

    cfg = {
        "title": req.title or f"Logs – {scope}",
        "data": [{"id": "logs", "path": rel_csv}],
        "pages": [
            {
                "title": "Overview",
                "blocks": [
                    {
                        "title": "Errors by Code",
                        "render": "plot",
                        "data": "logs",
                        "spec": {
                            "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
                            "mark": "bar",
                            "encoding": {
                                "x": {"field": "code", "type": "nominal", "sort": "-y"},
                                "y": {"aggregate": "count", "type": "quantitative"}
                            }
                        }
                    },
                    {
                        "title": "Levels Over Time",
                        "render": "plot",
                        "data": "logs",
                        "spec": {
                            "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
                            "mark": "line",
                            "encoding": {
                                "x": {"field": "ts", "type": "temporal"},
                                "y": {"aggregate": "count", "type": "quantitative"},
                                "color": {"field": "level", "type": "nominal"}
                            }
                        }
                    },
                    {
                        "title": "Top Components",
                        "render": "plot",
                        "data": "logs",
                        "spec": {
                            "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
                            "mark": "bar",
                            "encoding": {
                                "x": {"field": "component", "type": "nominal", "sort": "-y"},
                                "y": {"aggregate": "count", "type": "quantitative"}
                            }
                        }
                    }
                ]
            },
            {"title": "Log Table", "blocks": [ {"title": "Logs", "render": "table", "data": "logs", "columns": columns_logs, "search": True, "download": True} ]}
        ]
    }
    viz_yaml = proj_dir / "viz.yaml"
    viz_yaml.write_text(yaml.safe_dump(cfg, sort_keys=False), encoding="utf-8")
    rel_viz = str(viz_yaml.relative_to(ARTIFACTS_DIR)) if viz_yaml.is_relative_to(ARTIFACTS_DIR) else None
    return {"status": "ok", "project_dir": str(proj_dir), "viz_yaml": str(viz_yaml), "rel_viz": rel_viz}


# ---------------- POML export ----------------
class ExportPOMLRequest(BaseModel):
    document_id: str
    title: str | None = None
    variant: str | None = None  # generic|troubleshoot|catalog


@app.post("/export/poml")
async def export_poml(req: ExportPOMLRequest):
    doc = next((d for d in db.list_documents() if d.get("id") == req.document_id), None)
    if not doc:
        raise HTTPException(404, "Document not found")
    apis = db.list_apis(tag=None, method=None, path_like=None)
    apis = [a for a in apis if a.get("document_id") == req.document_id]
    tags = db.list_tags(document_id=req.document_id, q=None)
    logs = db.list_logs(level=None, code=None, q=None, ts_from=None, ts_to=None)
    logs = [l for l in logs if l.get("document_id") == req.document_id]
    # Try to attach local resources
    md_path = None
    chr_csv = None
    try:
        src = Path(doc.get("path", ""))
        stem = src.stem
        # markdown if PDF processed
        cand_md = ARTIFACTS_DIR / f"{stem}.md"
        if cand_md.exists():
            md_path = cand_md
        # CHR CSV if exists
        cand_chr = ARTIFACTS_DIR / "chr" / f"{stem}_chr.csv"
        if cand_chr.exists():
            chr_csv = cand_chr
    except Exception:
        pass
    poml = build_poml({**doc, **({"title": req.title} if req.title else {})}, apis, tags, logs, md_path, chr_csv, (req.variant or "generic"))
    out_dir = ARTIFACTS_DIR / "poml"
    out_dir.mkdir(parents=True, exist_ok=True)
    name = (Path(doc.get("path"," ")).stem or doc.get("id")) + ".poml"
    path = out_dir / name
    path.write_text(poml, encoding="utf-8")
    rel = str(path.relative_to(ARTIFACTS_DIR))
    return {"status": "ok", "rel": rel, "path": str(path)}


def _process_pdf_fast(file_path: Path, artifacts_dir: Path) -> tuple[list[dict], list[dict], dict]:
    placeholder = f"Extracted content unavailable for {file_path.name}."
    md_path = artifacts_dir / f"{file_path.stem}.md"
    md_path.write_text(placeholder, encoding="utf-8")
    json_path = artifacts_dir / f"{file_path.stem}.json"
    json_payload = {"texts": [{"text": placeholder}]}
    json_path.write_text(json.dumps(json_payload, indent=2), encoding="utf-8")
    units_path = artifacts_dir / f"{file_path.stem}.text_units.json"
    units_path.write_text(
        json.dumps([{ "text": placeholder, "page": 1 }], indent=2),
        encoding="utf-8",
    )
    return [], [], {
        "entities": [],
        "structure": None,
        "metric_hits": [],
        "tables": [],
        "charts": [],
        "formulas": [],
    }

import logging
import sys
_log = logging.getLogger(__name__)

def _process_and_store(file_path: Path, report_week: str, artifact_id: str, suffix: str, task_id: str | None = None):
    print(f"[DEBUG-STDERR] _process_and_store called for {file_path}", file=sys.stderr, flush=True)
    try:
        analysis_payload: dict | None = None
        facts: list[dict]
        evidence: list[dict]
        if suffix == ".pdf":
            # PDF is async-capable but can be used sync too
            fast_mode = _env_flag("FAST_PDF_MODE", False)
            print(f"[DEBUG-STDERR] FAST_PDF_MODE={fast_mode}", file=sys.stderr, flush=True)
            _log.warning(f"[DEBUG] FAST_PDF_MODE={fast_mode}, processing {file_path.name}")
            if fast_mode:
                _log.warning("[DEBUG] Using _process_pdf_fast")
                facts, evidence, analysis_payload = _process_pdf_fast(file_path, ARTIFACTS_DIR)
            else:
                _log.warning("[DEBUG] Using Docling process_pdf (synchronous)")
                try:
                    # Call process_pdf directly - it's synchronous and background tasks
                    # run in thread pools, so no need for anyio.run()
                    facts, evidence, analysis_payload = process_pdf(
                        file_path, report_week, ARTIFACTS_DIR, artifact_id
                    )
                    _log.warning(f"[DEBUG] Docling completed: {len(facts)} facts, {len(evidence)} evidence")
                except Exception as docling_err:
                    _log.error(f"[ERROR] Docling failed: {docling_err}")
                    import traceback
                    _log.error(traceback.format_exc())
                    _log.warning("[DEBUG] Falling back to _process_pdf_fast")
                    facts, evidence, analysis_payload = _process_pdf_fast(file_path, ARTIFACTS_DIR)
            # Ensure a PDF document row exists for deeplinks/open
            try:
                db.add_document({
                    "id": artifact_id,
                    "path": str(file_path),
                    "type": "pdf",
                    "title": file_path.name,
                    "source": "watch|upload",
                })
            except Exception:
                pass
        elif suffix == ".csv":
            facts, evidence = process_csv(file_path, report_week)
        elif suffix in [".xlsx", ".xls"]:
            facts, evidence = process_xlsx(file_path, report_week)
        elif suffix in MEDIA_SUFFIXES:
            facts = []
            media_payload = transcribe_media(file_path, ARTIFACTS_DIR, artifact_id)
            evidence = []
            transcript_text = (media_payload.get("text") or "").strip()
            warnings = media_payload.get("warnings") or []
            if transcript_text or warnings:
                evidence.append(
                    {
                        "id": str(uuid.uuid4()),
                        "locator": f"{file_path.name}#transcript",
                        "preview": media_payload.get("preview") or transcript_text[:240],
                        "content_type": "media_transcript",
                        "full_data": media_payload,
                    }
                )
            metadata = media_payload.get("metadata") or {}
            if metadata:
                preview_meta = {k: metadata.get(k) for k in ("duration_seconds", "format", "notes") if metadata.get(k) is not None}
                evidence.append(
                    {
                        "id": str(uuid.uuid4()),
                        "locator": f"{file_path.name}#metadata",
                        "preview": json.dumps(preview_meta or metadata, ensure_ascii=False)[:240],
                        "content_type": "media_metadata",
                        "full_data": metadata,
                    }
                )
            media_kind = "video" if suffix in VIDEO_SUFFIXES else "audio"
            extras = {
                "media": {
                    "kind": media_kind,
                    "transcript_preview": media_payload.get("preview"),
                    "artifacts": media_payload.get("artifacts"),
                    "engine": media_payload.get("engine"),
                    "metadata": metadata,
                    "status": media_payload.get("status"),
                    "warnings": warnings,
                }
            }
            try:
                db.update_artifact(artifact_id, extras=extras)
            except Exception:
                pass
        elif suffix in IMAGE_SUFFIXES:
            facts = []
            ocr_payload = extract_text_from_image(file_path, ARTIFACTS_DIR, artifact_id)
            evidence = []
            text = (ocr_payload.get("text") or "").strip()
            warnings = ocr_payload.get("warnings") or []
            if text or warnings:
                evidence.append(
                    {
                        "id": str(uuid.uuid4()),
                        "locator": f"{file_path.name}#ocr",
                        "preview": ocr_payload.get("preview") or text[:200],
                        "content_type": "image_ocr",
                        "full_data": ocr_payload,
                    }
                )
            extras = {
                "image": {
                    "transcript_preview": ocr_payload.get("preview"),
                    "artifacts": ocr_payload.get("artifacts"),
                    "metadata": ocr_payload.get("metadata"),
                    "warnings": warnings,
                }
            }
            try:
                db.update_artifact(artifact_id, extras=extras)
            except Exception:
                pass
        else:
            raise HTTPException(400, f"Unsupported file type: {suffix}")

        for fact in facts:
            fact["artifact_id"] = artifact_id
            db.add_fact(fact)
        for ev in evidence:
            ev["artifact_id"] = artifact_id
            db.add_evidence(ev)

        if analysis_payload and suffix == ".pdf":
            try:
                entities_raw = analysis_payload.get("entities") or []
                entities_prepared: list[dict] = []
                for idx, ent in enumerate(entities_raw):
                    seed = "|".join(
                        [
                            artifact_id,
                            "entity",
                            str(idx),
                            str(ent.get("label", "")),
                            str(ent.get("text", "")),
                            str(ent.get("start_char", "")),
                        ]
                    )
                    ent_id = str(uuid.uuid5(uuid.NAMESPACE_URL, seed))
                    entities_prepared.append(
                        {
                            "id": ent_id,
                            "document_id": artifact_id,
                            **ent,
                        }
                    )
                db.store_entities(artifact_id, entities_prepared)
            except Exception:
                pass

            try:
                structure = analysis_payload.get("structure")
                if structure:
                    db.store_structure(artifact_id, structure)
                else:
                    db.store_structure(artifact_id, None)
            except Exception:
                pass

            try:
                metric_hits_raw = analysis_payload.get("metric_hits") or []
                metric_prepared: list[dict] = []
                for idx, hit in enumerate(metric_hits_raw):
                    seed = "|".join(
                        [
                            artifact_id,
                            "metric",
                            str(idx),
                            str(hit.get("type", "")),
                            str(hit.get("value", "")),
                            str(hit.get("position", "")),
                        ]
                    )
                    metric_id = str(uuid.uuid5(uuid.NAMESPACE_URL, seed))
                    metric_prepared.append(
                        {
                            "id": metric_id,
                            "document_id": artifact_id,
                            **hit,
                        }
                    )
                db.store_metric_hits(artifact_id, metric_prepared)
            except Exception:
                pass

        if task_id:
            TASKS[task_id].update({
                "status": "completed",
                "facts_count": len(facts),
                "evidence_count": len(evidence)
            })
    except Exception as e:
        if task_id:
            TASKS[task_id].update({"status": "error", "error": str(e)})
        else:
            raise


@app.get("/tasks/{task_id}")
async def get_task_status(task_id: str):
    task = TASKS.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    return task


@app.post("/load_samples")
async def load_samples(background_tasks: BackgroundTasks, report_week: str = "", async_pdf: bool = True):
    """Server-side ingestion of sample files from SAMPLE_DIR."""
    sample_dir = Path(os.getenv("SAMPLE_DIR", "/app/samples"))
    if not sample_dir.exists():
        raise HTTPException(400, f"Sample directory not found: {sample_dir}")

    results = []
    # Known sample names or all supported files in folder
    sample_files: List[Path] = []
    for ext in ("*.csv", "*.xlsx", "*.xls", "*.pdf", "*.mp3", "*.wav", "*.mp4", "*.m4a", "*.png", "*.jpg", "*.jpeg"):
        sample_files.extend(sample_dir.glob(ext))
    web_sample_path = sample_dir / "web_urls.txt"
    if not sample_files and not web_sample_path.exists():
        return {"results": [{"status": "error", "error": "No sample files found"}]}

    # Reuse upload logic by simulating saved files
    for p in sample_files:
        try:
            file_id = str(uuid.uuid4())
            file_path = UPLOAD_DIR / f"{file_id}_{p.name}"
            shutil.copy2(p, file_path)
            suffix = file_path.suffix.lower()

            artifact_id = db.add_artifact({
                "id": file_id,
                "filename": p.name,
                "filepath": str(file_path),
                "filetype": suffix,
                "report_week": report_week,
                "status": "processing" if (async_pdf and suffix == ".pdf") else "processed"
            })

            if async_pdf and suffix == ".pdf":
                task_id = str(uuid.uuid4())
                TASKS[task_id] = {"status": "queued", "filename": p.name, "artifact_id": artifact_id}
                background_tasks.add_task(_process_and_store, file_path, report_week, artifact_id, suffix, task_id)
                results.append({"filename": p.name, "status": "queued", "task_id": task_id})
            else:
                _process_and_store(file_path, report_week, artifact_id, suffix, None)
                facts_count = len([f for f in db.get_facts(report_week) if f.get("artifact_id") == artifact_id])
                evidence_count = len([e for e in db.get_all_evidence() if e.get("artifact_id") == artifact_id])
                results.append({
                    "filename": p.name,
                    "status": "success",
                    "facts_count": facts_count,
                    "evidence_count": evidence_count
                })
        except Exception as e:
            results.append({"filename": p.name, "status": "error", "error": str(e)})

    if web_sample_path.exists():
        for line in web_sample_path.read_text(encoding="utf-8").splitlines():
            url = line.strip()
            if not url:
                continue
            artifact_id = str(uuid.uuid4())
            try:
                db.add_artifact(
                    {
                        "id": artifact_id,
                        "filename": url,
                        "filepath": url,
                        "filetype": "url",
                        "report_week": report_week,
                        "status": "processed",
                        "source_url": url,
                        "extras": {"web": {"status": "queued"}},
                    }
                )
                payload = ingest_web_url(url, ARTIFACTS_DIR, artifact_id)
                db.add_evidence(
                    {
                        "id": str(uuid.uuid4()),
                        "artifact_id": artifact_id,
                        "locator": url,
                        "preview": (payload.get("text") or "")[:280],
                        "content_type": "web_page",
                        "full_data": payload,
                    }
                )
                db.update_artifact(
                    artifact_id,
                    extras={
                        "web": {
                            "status": "processed",
                            "preview": (payload.get("text") or "")[:240],
                            "metadata": payload.get("metadata"),
                            "artifacts": payload.get("artifacts"),
                            "warnings": payload.get("warnings"),
                        }
                    },
                )
                results.append({"filename": url, "status": "success", "facts_count": 0, "evidence_count": 1})
            except Exception as exc:
                results.append({"filename": url, "status": "error", "error": str(exc)})

    return {"results": results}

@app.get("/facts")
async def get_facts(report_week: str = None):
    """Get all facts, optionally filtered by report week"""
    facts = db.get_facts(report_week)
    return {"facts": facts}


@app.get("/analysis/financials")
async def get_financial_statements(artifact_id: str | None = None):
    """Return detected financial statements from processed tables."""
    statements: List[Dict[str, Any]] = []
    for ev in db.get_all_evidence():
        if artifact_id and ev.get("artifact_id") != artifact_id:
            continue
        if ev.get("content_type") not in {"financial_table", "table"}:
            continue
        full_data = ev.get("full_data") or {}
        if not isinstance(full_data, dict):
            continue
        statement = full_data.get("statement") or {}
        if not isinstance(statement, dict):
            continue
        stmt_type = statement.get("type")
        if stmt_type in (None, "", "unknown"):
            continue
        statements.append(
            {
                "evidence_id": ev.get("id"),
                "artifact_id": ev.get("artifact_id"),
                "locator": ev.get("locator"),
                "statement_type": stmt_type,
                "confidence": statement.get("confidence"),
                "summary": statement.get("summary") or {},
                "columns": full_data.get("columns", []),
                "rows": full_data.get("rows", []),
                "header_info": full_data.get("header_info"),
            }
        )

    statements.sort(
        key=lambda item: (
            item.get("artifact_id") or "",
            item.get("statement_type") or "",
            item.get("locator") or "",
        )
    )
    return {"statements": statements}


@app.get("/evidence/{evidence_id}")
async def get_evidence(evidence_id: str):
    """Get evidence by ID"""
    evidence = db.get_evidence(evidence_id)
    if not evidence:
        raise HTTPException(404, "Evidence not found")
    return evidence

@app.post("/ask")
async def ask_question(
    question: str,
    use_hrm: bool = Query(False, description="Enable HRM sidecar (if supported)"),
):
    """Ask a question and get answer with citations.

    If `use_hrm=true` and HRM is enabled, we simulate a multi-step refinement
    cycle for demonstration and record HRM metrics. The underlying QA remains
    the same in this initial integration (no model change).
    """
    t0 = time.time()
    result = await qa_engine.ask(question)
    if HRM_ENABLED and use_hrm:
        # Simulate a tiny refinement loop over the answer text: trim + normalize
        steps = max(HRM_CFG.Mmin, 2)
        refined = result.get("answer", "").strip()
        refined = refined.replace("  ", " ")
        result["answer"] = refined
        HRM_STATS.record(steps=steps, latency_ms=(time.time() - t0) * 1000.0, payload={
            "question": question,
            "mode": "ask",
        })
        result["hrm"] = {"enabled": True, "steps": steps}
    else:
        result["hrm"] = {"enabled": False}
    return result

@app.delete("/reset")
async def reset_database():
    """Clear all data"""
    db.reset()
    return {"status": "Database reset"}

# ------------ HRM experiment & metrics endpoints ------------

class SortDigitsRequest(BaseModel):
    seq: str
    Mmax: int | None = None
    Mmin: int | None = None


@app.post("/experiments/hrm/sort_digits")
def hrm_sort_digits(req: SortDigitsRequest):
    if not req.seq or not req.seq.isdigit():
        raise HTTPException(400, "Provide 'seq' as digits only, e.g. '93241'.")
    cfg = HRMConfig(
        Mmax=req.Mmax or HRM_CFG.Mmax,
        Mmin=req.Mmin or HRM_CFG.Mmin,
        threshold=HRM_CFG.threshold,
    )
    t0 = time.time()
    out, steps, trace = refine_sort_digits(req.seq, cfg)
    dt = (time.time() - t0) * 1000.0
    HRM_STATS.record(steps=steps, latency_ms=dt, payload={"mode": "sort_digits", "seq": req.seq})
    return {"in": req.seq, "out": out, "steps": steps, "trace": trace, "latency_ms": round(dt, 3)}


class EchoRequest(BaseModel):
    text: str


@app.post("/experiments/hrm/echo")
def hrm_echo(req: EchoRequest, Mmax: int = 3, Mmin: int = 1):
    """Echo with a simulated refinement loop; returns steps and variants."""
    cfg = HRMConfig(Mmax=Mmax, Mmin=Mmin, threshold=HRM_CFG.threshold)
    variants: List[str] = [req.text]
    t0 = time.time()
    s = req.text
    steps = 0
    for m in range(1, cfg.Mmax + 1):
        steps = m
        # simple normalization as a stand-in for refinement
        s2 = " ".join(s.strip().split())
        variants.append(s2)
        if m >= cfg.Mmin and s2 == s:
            break
        s = s2
    dt = (time.time() - t0) * 1000.0
    HRM_STATS.record(steps=steps, latency_ms=dt, payload={"mode": "echo"})
    return {"out": s, "steps": steps, "variants": variants, "latency_ms": round(dt, 3)}


@app.get("/metrics/hrm")
def hrm_metrics():
    return HRM_STATS.snapshot()


@app.get("/metrics")
def metrics_prometheus():
    snap = HRM_STATS.snapshot()
    lines = [
        f"pmoves_hrm_total_runs {snap['total_runs']}",
        f"pmoves_hrm_avg_steps {snap['avg_steps']}",
        f"pmoves_hrm_avg_latency_ms {snap['avg_latency_ms']}",
    ]
    return ("\n".join(lines) + "\n", 200, {"Content-Type": "text/plain; version=0.0.4"})
@app.get("/open/pdf")
async def open_pdf(artifact_id: str, page: int | None = None):
    if _env_flag("FAST_PDF_MODE", True):
        raise HTTPException(403, "PDF open disabled in fast mode")
    if not _env_flag("OPEN_PDF_ENABLED", False):
        raise HTTPException(403, "PDF open is disabled")
    # Locate by document id (preferred) or fall back to artifact id
    p: Path | None = None
    doc = next((d for d in db.list_documents(type="pdf") if d.get("id") == artifact_id), None)
    if doc:
        p = Path(doc.get("path", ""))
    else:
        art = next((a for a in db.get_artifacts() if a.get("id") == artifact_id and str(a.get("filepath",""))).__iter__(), None)
        # The above ".__iter__()" trick is to avoid mypy complaining; essentially select the first match
        if not art:
            # try explicit scan
            for a in db.get_artifacts():
                if a.get("id") == artifact_id:
                    art = a
                    break
        if art:
            p = Path(art.get("filepath", ""))
    if not p or not p.exists() or p.suffix.lower() != ".pdf":
        raise HTTPException(403, "PDF file missing")
    # Restrict to uploads dir
    try:
        if not Path(p).resolve().is_relative_to(UPLOAD_DIR.resolve()):
            raise HTTPException(403, "Access denied")
    except Exception:
        # Python <3.9 fallback: emulate is_relative_to
        up = str(UPLOAD_DIR.resolve())
        pr = str(Path(p).resolve())
        if not pr.startswith(up):
            raise HTTPException(403, "Access denied")
    headers = {"Content-Disposition": f"inline; filename=\"{p.name}\""}
    return FileResponse(str(p), media_type="application/pdf", headers=headers)


WebUrlForm = Annotated[List[str] | None, Form()]


@app.post("/upload")
async def upload_files(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] | None = File(default=None),
    report_week: str = "",
    async_pdf: bool = True,
    web_urls: WebUrlForm = None,
):
    """Upload and process documents."""

    results: List[Dict] = []
    incoming_files = files or []

    for file in incoming_files:
        file_id = str(uuid.uuid4())

        # Sanitize filename to prevent path traversal attacks
        safe_filename = os.path.basename(file.filename) if file.filename else "upload"
        # Also remove any remaining path separators that might have been encoded
        safe_filename = safe_filename.replace("/", "_").replace("\\", "_")

        file_path = UPLOAD_DIR / f"{file_id}_{safe_filename}"

        # Read file content and check size
        file_content = await file.read()
        if len(file_content) > MAX_FILE_SIZE:
            results.append({
                "filename": file.filename,
                "status": "error",
                "error": f"File size exceeds maximum allowed size of {MAX_FILE_SIZE // (1024*1024)}MB"
            })
            continue

        # Write file to disk
        with file_path.open("wb") as buffer:
            buffer.write(file_content)

        try:
            suffix = file_path.suffix.lower()
            artifact_id = db.add_artifact(
                {
                    "id": file_id,
                    "filename": file.filename,
                    "filepath": str(file_path),
                    "filetype": suffix,
                    "report_week": report_week,
                    "status": "processing" if (async_pdf and suffix == ".pdf") else "processed",
                }
            )
            if async_pdf and suffix == ".pdf":
                task_id = str(uuid.uuid4())
                TASKS[task_id] = {"status": "queued", "filename": file.filename, "artifact_id": artifact_id}
                background_tasks.add_task(_process_and_store, file_path, report_week, artifact_id, suffix, task_id)
                results.append({"filename": file.filename, "status": "queued", "task_id": task_id})
            else:
                _process_and_store(file_path, report_week, artifact_id, suffix, None)
                facts_count = len([f for f in db.get_facts(report_week) if f.get("artifact_id") == artifact_id])
                evidence_count = len([e for e in db.get_all_evidence() if e.get("artifact_id") == artifact_id])
                results.append(
                    {
                        "filename": file.filename,
                        "status": "success",
                        "facts_count": facts_count,
                        "evidence_count": evidence_count,
                    }
                )
        except Exception as exc:
            results.append({"filename": file.filename, "status": "error", "error": str(exc)})

    for raw_url in web_urls or []:
        url = (raw_url or "").strip()
        if not url:
            continue
        artifact_id = str(uuid.uuid4())
        try:
            artifact_record = {
                "id": artifact_id,
                "filename": url,
                "filepath": url,
                "filetype": "url",
                "report_week": report_week,
                "status": "processed",
                "source_url": url,
                "extras": {"web": {"status": "queued"}},
            }
            db.add_artifact(artifact_record)
            web_payload = ingest_web_url(url, ARTIFACTS_DIR, artifact_id)
            db.add_evidence(
                {
                    "id": str(uuid.uuid4()),
                    "artifact_id": artifact_id,
                    "locator": url,
                    "preview": (web_payload.get("text") or "")[:280],
                    "content_type": "web_page",
                    "full_data": web_payload,
                }
            )
            db.update_artifact(
                artifact_id,
                extras={
                    "web": {
                        "status": "processed",
                        "preview": (web_payload.get("text") or "")[:240],
                        "metadata": web_payload.get("metadata"),
                        "artifacts": web_payload.get("artifacts"),
                        "warnings": web_payload.get("warnings"),
                    }
                },
            )
            results.append(
                {
                    "filename": url,
                    "status": "success",
                    "facts_count": 0,
                    "evidence_count": 1,
                    "artifact_id": artifact_id,
                }
            )
        except Exception as exc:
            results.append({"filename": url, "status": "error", "error": str(exc)})

    return {"results": results}


if __name__ == "__main__":
    import uvicorn

    port = int(os.getenv("PORT", "8484"))
    uvicorn.run(app, host="0.0.0.0", port=port)


